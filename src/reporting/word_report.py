# src/reporting/word_report.py

import pandas as pd
import re
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENTATION
from pathlib import Path
from src.utils.exceptions import ReportingError
from src.utils.logging import setup_logger

logger = setup_logger(__name__)

# Nombres de meses en español
MESES = {
    1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
    5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
    9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
}


def _extract_year_month_from_path(path: Path):
    """
    Extrae el año y mes del nombre del archivo.
    El formato esperado es: YYYYMMreporte.docx
    """
    filename = path.stem  # Obtiene el nombre sin extensión
    match = re.match(r'(\d{4})(\d{2})reporte', filename)
    if match:
        year = int(match.group(1))
        month = int(match.group(2))
        return year, month
    raise ReportingError(f"No se pudo extraer año y mes del nombre del archivo: {filename}")


def _calculate_tramo_statistics(df: pd.DataFrame):
    """
    Calcula las estadísticas por tramos de saldo.
    
    Retorna tres diccionarios:
    - cred: Créditos por tramo
    - deb: Débitos por tramo
    - sald: Saldos y cantidad de socios por tramo
    """
    # Filtrar solo registros con saldo distinto de cero
    capital = df[df['Saldo'] != 0].copy()
    
    # Tramo 1: Saldo <= 10,000
    tramo1 = capital[capital['Saldo'] <= 10000]
    # Tramo 2: 10,001 < Saldo <= 50,000
    tramo2 = capital[(capital['Saldo'] > 10000) & (capital['Saldo'] <= 50000)]
    # Tramo 3: 50,001 < Saldo <= 100,000
    tramo3 = capital[(capital['Saldo'] > 50000) & (capital['Saldo'] <= 100000)]
    # Tramo 4: Saldo > 100,000
    tramo4 = capital[capital['Saldo'] > 100000]
    
    # Calcular créditos por tramo
    cred = {
        '1': int(tramo1['Credito'].sum()),
        '2': int(tramo2['Credito'].sum()),
        '3': int(tramo3['Credito'].sum()),
        '4': int(tramo4['Credito'].sum()),
        'total': int(capital['Credito'].sum())
    }
    
    # Calcular débitos por tramo
    deb = {
        '1': int(tramo1['Debito'].sum()),
        '2': int(tramo2['Debito'].sum()),
        '3': int(tramo3['Debito'].sum()),
        '4': int(tramo4['Debito'].sum()),
        'total': int(capital['Debito'].sum())
    }
    
    # Calcular saldos y cantidad de socios por tramo
    sald = {
        '1': [int(tramo1['Saldo'].sum()), len(tramo1)],
        '2': [int(tramo2['Saldo'].sum()), len(tramo2)],
        '3': [int(tramo3['Saldo'].sum()), len(tramo3)],
        '4': [int(tramo4['Saldo'].sum()), len(tramo4)],
        'total': int(capital['Saldo'].sum())
    }
    
    # Validar que las sumas coincidan
    assert cred['1'] + cred['2'] + cred['3'] + cred['4'] == cred['total']
    assert deb['1'] + deb['2'] + deb['3'] + deb['4'] == deb['total']
    assert sald['1'][0] + sald['2'][0] + sald['3'][0] + sald['4'][0] == sald['total']
    assert cred['total'] - deb['total'] == sald['total']
    
    return cred, deb, sald


def _format_number(value: int) -> str:
    """Formatea un número con puntos como separadores de miles."""
    return '{:,}'.format(value).replace(',', '.')


def generate_word_report(df: pd.DataFrame, output_path: Path):
    """
    Genera un reporte Word con el formato del documento "Capital Pagado".
    
    El documento incluye:
    - Título y mes/año
    - Datos del balance tributario (Total Créditos, Total Débitos, Saldo Acreedor)
    - Clasificación de aportes por socio con 4 tramos de saldo
    - Formato landscape (paisaje)
    - Fuente Calibri 14pt
    
    Args:
        df: DataFrame con columnas: Rut, Debito, Credito, Saldo, Cuotas, Nombre
        output_path: Ruta donde se guardará el archivo Word
    """
    logger.info(f"Generando reporte Word: {output_path.name}")

    try:
        # Asegurar que el directorio existe
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Validar columnas requeridas
        required_cols = ['Rut', 'Debito', 'Credito', 'Saldo', 'Nombre']
        missing_cols = [col for col in required_cols if col not in df.columns]
        if missing_cols:
            raise ReportingError(
                f"El DataFrame no tiene las columnas requeridas: {missing_cols}"
            )
        
        # Extraer año y mes del nombre del archivo
        year, month = _extract_year_month_from_path(output_path)
        mes_nombre = MESES.get(month, f"Mes {month}")
        
        # Calcular estadísticas por tramos
        cred, deb, sald = _calculate_tramo_statistics(df)
        
        # Datos para el documento
        total_creditos = cred['total']
        total_debitos = deb['total']
        saldo_acreedor = sald['total']
        total_socios = sald['1'][1] + sald['2'][1] + sald['3'][1] + sald['4'][1]
        
        personas_tramo1 = sald['1'][1]
        personas_tramo2 = sald['2'][1]
        personas_tramo3 = sald['3'][1]
        personas_tramo4 = sald['4'][1]
        
        # Crear documento
        informe = Document()
        
        # Configuración: orientación landscape y fuente
        section = informe.sections[-1]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        
        # Configurar fuente Calibri 14pt
        style = informe.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(14)
        
        # Escribir contenido del documento
        informe.add_paragraph("\t\tCuenta\t\t\t\t\tCapital Pagado.")
        
        informe.add_paragraph(f"\t\tMes\t\t\t\t\t\t{mes_nombre} {year}")
        
        # Datos del balance tributario
        datos_tributarios = "\t\tDatos del Balance tributario:"
        datos_tributarios += f"\n\t\tTotal Creditos\t\t\t\t${_format_number(total_creditos)}"
        datos_tributarios += f"\n\t\tTotal Debitos\t\t\t\t$  {_format_number(total_debitos)}"
        datos_tributarios += f"\n\t\tSaldo Acreedor\t\t\t\t${_format_number(saldo_acreedor)}"
        datos_tributarios += "\n\t\t\t\t\t\t\t\t============"
        informe.add_paragraph(datos_tributarios)
        
        # Clasificación aportes por socio
        informe.add_paragraph("Clasificación aportes por socio")
        informe.add_paragraph("Desde\t\t\tHasta\t\tN.º socios\t\tDébitos\t\tCréditos\t\tSaldo")
        
        # Tramo 1: 0 - 10,000
        doc_tramo1 = f"{personas_tramo1}\t\t${_format_number(deb['1'])}"
        doc_tramo1 += f"\t\t${_format_number(cred['1'])}"
        doc_tramo1 += f"\t\t${_format_number(sald['1'][0])}"
        informe.add_paragraph(f"\t0\t\t10.000\t\t{doc_tramo1}")
        
        # Tramo 2: 10,001 - 50,000
        doc_tramo2 = f"{personas_tramo2}\t\t${_format_number(deb['2'])}"
        doc_tramo2 += f"\t\t${_format_number(cred['2'])}"
        doc_tramo2 += f"\t\t${_format_number(sald['2'][0])}"
        informe.add_paragraph(f"10.001\t\t50.000\t\t{doc_tramo2}")
        
        # Tramo 3: 50,001 - 100,000
        doc_tramo3 = f"{personas_tramo3}\t\t${_format_number(deb['3'])}"
        doc_tramo3 += f"\t\t${_format_number(cred['3'])}"
        doc_tramo3 += f"\t\t${_format_number(sald['3'][0])}"
        informe.add_paragraph(f"50.001\t\t100.000\t\t{doc_tramo3}")
        
        # Tramo 4: 100,001 o superior
        doc_tramo4 = f"{personas_tramo4}\t\t${_format_number(deb['4'])}"
        doc_tramo4 += f"\t\t${_format_number(cred['4'])}"
        doc_tramo4 += f"\t${_format_number(sald['4'][0])}"
        informe.add_paragraph(f"100.001\to superior\t\t\t{doc_tramo4}")
        
        # Separador
        informe.add_paragraph(
            "\t\t\t\t\t-----------------\t---------------------------------------------------------------------"
        )
        
        # Totales
        doc_totales = f"Totales\t\t\t\t\t{total_socios}"
        doc_totales += f"\t\t${_format_number(total_debitos)}"
        doc_totales += f"\t\t${_format_number(total_creditos)}"
        doc_totales += f"\t${_format_number(saldo_acreedor)}"
        doc_totales += "\n\t\t\t\t\t===========\t=========================================="
        informe.add_paragraph(doc_totales)
        
        # Guardar documento
        informe.save(output_path)
        
        logger.info(f"Reporte Word generado exitosamente: {output_path}")
        
    except ReportingError:
        raise
    except Exception as e:
        logger.exception("Error generando reporte Word")
        raise ReportingError("No se pudo generar el reporte Word") from e
